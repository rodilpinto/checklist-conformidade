# -*- coding: utf-8 -*-
"""
Modulo de extracao de texto para o app Streamlit de checklist.

Suporta extracao de conteudo textual a partir de PDFs, DOCX, URLs e texto puro.
Todas as funcoes recebem bytes ou strings e devolvem texto limpo em UTF-8.

Dependencias externas:
    - pymupdf (import fitz)
    - python-docx (from docx import Document)
    - requests
    - beautifulsoup4 (from bs4 import BeautifulSoup)
"""

from __future__ import annotations

import io
import ipaddress
import logging
import socket
from typing import Union
from urllib.parse import urlparse

import fitz  # pymupdf
import requests
from bs4 import BeautifulSoup
from docx import Document

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------

_PAGE_SEPARATOR = "\n\n--- Pagina {page_num} ---\n\n"

_DEFAULT_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

_REQUEST_TIMEOUT_SECONDS = 30

# Limite maximo de tamanho de entrada para evitar ataques de file-bomb.
# 50 MB para PDFs/DOCX (suficiente para normativos governamentais),
# 10 MB para respostas HTTP.
# SECURITY: 200 MB era permissivo demais -- PDFs comprimidos podem expandir
# 10-50x em memoria durante a extracao de texto (zip-bomb pattern).
_MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024   # 50 MB
_MAX_RESPONSE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

# Limite de caracteres do texto extraido antes de enviar ao LLM.
# Previne que arquivos grandes (mesmo dentro do limite de bytes) gerem
# textos excessivamente longos que consumam memoria e tokens da API.
_MAX_EXTRACTED_CHARS = 500_000  # ~375k tokens estimados

# Tags HTML cujo conteudo deve ser removido antes da extracao de texto.
_TAGS_TO_STRIP = ["script", "style", "noscript", "svg", "canvas"]

# Dominios/IPs bloqueados para prevenir SSRF.
# Requisicoes a enderecos privados, loopback e link-local sao recusadas.
_BLOCKED_HOSTNAMES = {"localhost", "metadata.google.internal"}


# ---------------------------------------------------------------------------
# Helpers de seguranca
# ---------------------------------------------------------------------------


def _is_private_ip(hostname: str) -> bool:
    """Verifica se um hostname resolve para um IP privado, loopback ou link-local.

    Previne ataques SSRF que tentam acessar servicos internos (ex.: cloud
    metadata em 169.254.169.254) atraves do extrator de URL.

    Args:
        hostname: Nome de host ou endereco IP a ser verificado.

    Returns:
        True se o endereco for privado/reservado, False caso contrario.
    """
    try:
        # Tenta interpretar diretamente como IP (ex.: "169.254.169.254")
        addr = ipaddress.ip_address(hostname)
        return addr.is_private or addr.is_loopback or addr.is_link_local or addr.is_reserved
    except ValueError:
        pass

    # Resolve hostname para IP(s) e verifica cada um
    try:
        resolved = socket.getaddrinfo(hostname, None)
    except socket.gaierror:
        # Se nao consegue resolver, bloqueia por precaucao
        return True

    for family, _type, _proto, _canonname, sockaddr in resolved:
        ip_str = sockaddr[0]
        try:
            addr = ipaddress.ip_address(ip_str)
            if addr.is_private or addr.is_loopback or addr.is_link_local or addr.is_reserved:
                return True
        except ValueError:
            # IP malformado -- bloqueia por precaucao
            return True

    return False


def _validate_url(url: str) -> str:
    """Valida e sanitiza uma URL antes de realizar a requisicao.

    Verifica esquema (http/https), presenca de hostname, e bloqueia
    enderecos internos/privados para prevenir SSRF.

    Args:
        url: URL a ser validada.

    Returns:
        A URL limpa (stripped) se valida.

    Raises:
        ValueError: Se a URL for invalida ou apontar para endereco bloqueado.
    """
    if not url or not isinstance(url, str):
        raise ValueError("A URL nao foi fornecida ou e invalida.")

    url = url.strip()

    if not url.startswith(("http://", "https://")):
        raise ValueError(
            f"A URL deve iniciar com 'http://' ou 'https://'. Recebido: '{url}'"
        )

    parsed = urlparse(url)
    hostname = parsed.hostname

    if not hostname:
        raise ValueError(f"Nao foi possivel extrair o hostname da URL: '{url}'")

    # Bloqueia hostnames conhecidos como perigosos
    if hostname.lower() in _BLOCKED_HOSTNAMES:
        raise ValueError(
            f"Acesso bloqueado: o hostname '{hostname}' nao e permitido."
        )

    # Bloqueia IPs privados / loopback / link-local (prevencao SSRF)
    if _is_private_ip(hostname):
        raise ValueError(
            f"Acesso bloqueado: o endereco '{hostname}' resolve para um IP "
            f"privado ou reservado."
        )

    return url


def _check_file_size(file_bytes: bytes, label: str) -> None:
    """Valida que os bytes de entrada nao excedem o limite de seguranca.

    Args:
        file_bytes: Conteudo binario do arquivo.
        label: Rotulo descritivo do tipo de arquivo (ex.: "PDF", "DOCX").

    Raises:
        ValueError: Se o arquivo exceder ``_MAX_FILE_SIZE_BYTES``.
    """
    if len(file_bytes) > _MAX_FILE_SIZE_BYTES:
        size_mb = len(file_bytes) / (1024 * 1024)
        limit_mb = _MAX_FILE_SIZE_BYTES / (1024 * 1024)
        raise ValueError(
            f"O arquivo {label} tem {size_mb:.1f} MB, que excede o limite "
            f"de {limit_mb:.0f} MB."
        )


def _truncate_extracted_text(text: str, label: str) -> str:
    """Trunca o texto extraido ao limite maximo de caracteres.

    Previne que arquivos grandes (mesmo dentro do limite de bytes) gerem
    textos que causem problemas de memoria ou excedam os limites da API.
    Loga um aviso se o truncamento for aplicado.

    SECURITY: Mitiga expansao de arquivo comprimido (zip-bomb pattern) onde
    poucos bytes no disco geram megabytes de texto em memoria.

    Args:
        text: Texto extraido da fonte.
        label: Nome da fonte para mensagem de log (ex.: "PDF", "DOCX").

    Returns:
        Texto original se dentro do limite, ou texto truncado com aviso.
    """
    if len(text) <= _MAX_EXTRACTED_CHARS:
        return text

    logger.warning(
        "Texto extraido de %s truncado de %d para %d caracteres (limite de seguranca).",
        label,
        len(text),
        _MAX_EXTRACTED_CHARS,
    )
    # Trunca no limite de caracteres preservando palavras inteiras quando possivel
    truncated = text[:_MAX_EXTRACTED_CHARS]
    last_space = truncated.rfind(" ")
    if last_space > _MAX_EXTRACTED_CHARS * 0.95:  # trunca na ultima palavra se proximo do limite
        truncated = truncated[:last_space]
    return truncated


# ---------------------------------------------------------------------------
# Funcoes de extracao por tipo de fonte
# ---------------------------------------------------------------------------


def extract_from_pdf(file_bytes: bytes) -> str:
    """Extrai texto de todas as paginas de um arquivo PDF.

    Utiliza pymupdf (fitz) para abrir o PDF diretamente a partir dos bytes
    recebidos -- compativel com ``st.file_uploader().read()``.

    Args:
        file_bytes: Conteudo binario do arquivo PDF.

    Returns:
        Texto concatenado de todas as paginas, separado por marcadores
        indicando o numero de cada pagina.

    Raises:
        ValueError: Se ``file_bytes`` estiver vazio, for None ou exceder
            o tamanho maximo permitido.
        RuntimeError: Se a extracao falhar (PDF corrompido, protegido, etc.).
    """
    if not file_bytes:
        raise ValueError("Os bytes do PDF estao vazios ou nao foram fornecidos.")

    _check_file_size(file_bytes, "PDF")

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise RuntimeError(
            f"Nao foi possivel abrir o PDF. Verifique se o arquivo e valido. "
            f"Detalhe: {exc}"
        ) from exc

    parts: list[str] = []
    try:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if text and text.strip():
                parts.append(
                    _PAGE_SEPARATOR.format(page_num=page_num) + text.strip()
                )
    finally:
        doc.close()

    if not parts:
        logger.warning("O PDF foi processado mas nenhum texto foi extraido.")
        return ""

    return "\n".join(parts).strip()


def extract_from_docx(file_bytes: bytes) -> str:
    """Extrai texto de paragrafos e tabelas de um arquivo DOCX.

    Utiliza python-docx para abrir o documento a partir de um buffer
    ``BytesIO`` -- compativel com ``st.file_uploader().read()``.

    Alem dos paragrafos do corpo, extrai tambem o conteudo de todas as
    tabelas do documento, pois documentos governamentais frequentemente
    organizam informacoes em formato tabular.

    Args:
        file_bytes: Conteudo binario do arquivo DOCX.

    Returns:
        Texto de todos os paragrafos e celulas de tabela, concatenados.

    Raises:
        ValueError: Se ``file_bytes`` estiver vazio, for None ou exceder
            o tamanho maximo permitido.
        RuntimeError: Se a extracao falhar (arquivo corrompido, formato
            invalido, etc.).
    """
    if not file_bytes:
        raise ValueError("Os bytes do DOCX estao vazios ou nao foram fornecidos.")

    _check_file_size(file_bytes, "DOCX")

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise RuntimeError(
            f"Nao foi possivel abrir o DOCX. Verifique se o arquivo e valido. "
            f"Detalhe: {exc}"
        ) from exc

    # Filtra paragrafos vazios para evitar linhas em branco excessivas
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    # Extrai conteudo de tabelas (comum em documentos governamentais)
    table_texts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_cells:
                table_texts.append(" | ".join(row_cells))

    all_parts = paragraphs + table_texts

    if not all_parts:
        logger.warning("O DOCX foi processado mas nenhum texto foi extraido.")
        return ""

    return "\n".join(all_parts)


def extract_from_url(url: str) -> str:
    """Extrai texto limpo do corpo de uma pagina web.

    Realiza uma requisicao GET com timeout e User-Agent de navegador,
    remove tags de script/style e retorna apenas o conteudo textual.

    Inclui protecao contra SSRF: URLs que apontem para enderecos privados,
    loopback ou link-local sao bloqueadas antes da requisicao.

    Args:
        url: Endereco completo da pagina (deve iniciar com http:// ou https://).

    Returns:
        Texto limpo extraido do ``<body>`` da pagina.

    Raises:
        ValueError: Se a URL estiver vazia, com formato invalido ou apontar
            para endereco bloqueado.
        RuntimeError: Se a requisicao falhar (timeout, erro HTTP, etc.) ou
            se a resposta exceder o tamanho maximo permitido.
    """
    url = _validate_url(url)

    headers = {"User-Agent": _DEFAULT_USER_AGENT}

    try:
        # SECURITY: allow_redirects=False na primeira requisicao.
        # Redirecionamentos sao seguidos manualmente apos revalidar o destino,
        # prevenindo bypass de SSRF via redirect (ex.: URL publica -> 302 ->
        # http://169.254.169.254/latest/meta-data/).
        response = requests.get(
            url,
            headers=headers,
            timeout=_REQUEST_TIMEOUT_SECONDS,
            stream=True,
            allow_redirects=False,
        )

        # Segue redirecionamentos manualmente, revalidando cada destino
        redirect_count = 0
        _MAX_REDIRECTS = 5
        while response.is_redirect and redirect_count < _MAX_REDIRECTS:
            redirect_url = response.headers.get("Location", "")
            if not redirect_url:
                break
            # URLs relativas precisam ser resolvidas antes da validacao
            if redirect_url.startswith("/"):
                parsed_orig = urlparse(url)
                redirect_url = f"{parsed_orig.scheme}://{parsed_orig.netloc}{redirect_url}"
            # Revalida o destino do redirecionamento contra SSRF
            redirect_url = _validate_url(redirect_url)
            response.close()
            response = requests.get(
                redirect_url,
                headers=headers,
                timeout=_REQUEST_TIMEOUT_SECONDS,
                stream=True,
                allow_redirects=False,
            )
            url = redirect_url
            redirect_count += 1

        response.raise_for_status()

        # Verifica tamanho via Content-Length antes de baixar o corpo inteiro
        content_length = response.headers.get("Content-Length")
        if content_length and int(content_length) > _MAX_RESPONSE_SIZE_BYTES:
            response.close()
            size_mb = int(content_length) / (1024 * 1024)
            limit_mb = _MAX_RESPONSE_SIZE_BYTES / (1024 * 1024)
            raise RuntimeError(
                f"A pagina '{url}' reporta {size_mb:.1f} MB, excedendo o "
                f"limite de {limit_mb:.0f} MB."
            )

        # Le o conteudo com limite de tamanho (stream=True evita download imediato)
        chunks: list[bytes] = []
        downloaded = 0
        for chunk in response.iter_content(chunk_size=8192):
            downloaded += len(chunk)
            if downloaded > _MAX_RESPONSE_SIZE_BYTES:
                response.close()
                limit_mb = _MAX_RESPONSE_SIZE_BYTES / (1024 * 1024)
                raise RuntimeError(
                    f"A resposta de '{url}' excedeu o limite de "
                    f"{limit_mb:.0f} MB durante o download."
                )
            chunks.append(chunk)

        raw_content = b"".join(chunks)

    except requests.exceptions.Timeout as exc:
        raise RuntimeError(
            f"A requisicao para '{url}' excedeu o tempo limite "
            f"de {_REQUEST_TIMEOUT_SECONDS} segundos."
        ) from exc
    except requests.exceptions.ConnectionError as exc:
        raise RuntimeError(
            f"Nao foi possivel conectar a '{url}'. Verifique a URL e sua conexao."
        ) from exc
    except requests.exceptions.HTTPError as exc:
        raise RuntimeError(
            f"A pagina retornou um erro HTTP: {exc.response.status_code}. "
            f"URL: '{url}'"
        ) from exc
    except requests.exceptions.RequestException as exc:
        raise RuntimeError(
            f"Erro ao acessar '{url}'. Detalhe: {exc}"
        ) from exc

    # Detecta encoding a partir dos headers ou do conteudo.
    # Nao usar response.apparent_encoding pois o stream ja foi consumido.
    encoding = response.encoding
    if not encoding or encoding == "ISO-8859-1":
        # Fallback: tentar detectar via chardet se disponivel, senao utf-8
        try:
            import chardet
            detected = chardet.detect(raw_content[:10_000])
            encoding = detected.get("encoding") or "utf-8"
        except ImportError:
            encoding = "utf-8"
    text_content = raw_content.decode(encoding, errors="replace")

    soup = BeautifulSoup(text_content, "html.parser")

    # Remove tags que nao contribuem conteudo textual
    for tag in soup.find_all(_TAGS_TO_STRIP):
        tag.decompose()

    body = soup.find("body")
    if body is None:
        # Fallback: usa o documento inteiro se nao houver <body>
        body = soup

    text = body.get_text(separator="\n", strip=True)

    if not text:
        logger.warning("A pagina '%s' foi processada mas nenhum texto foi extraido.", url)
        return ""

    return text


# ---------------------------------------------------------------------------
# Dispatcher unificado
# ---------------------------------------------------------------------------


def extract_text(source: Union[bytes, str], source_type: str) -> str:
    """Dispatcher unificado para extracao de texto de diferentes fontes.

    Seleciona automaticamente a funcao de extracao adequada com base no
    ``source_type`` informado.

    Args:
        source: Dados da fonte. Pode ser ``bytes`` (para PDF/DOCX),
            ``str`` (para URL ou texto puro).
        source_type: Tipo da fonte. Valores aceitos:
            - ``"pdf"``  -- arquivo PDF (source deve ser bytes)
            - ``"docx"`` -- arquivo DOCX (source deve ser bytes)
            - ``"url"``  -- endereco web (source deve ser str)
            - ``"text"`` -- texto puro (source deve ser str, retornado diretamente)

    Returns:
        Texto extraido da fonte.

    Raises:
        ValueError: Se ``source_type`` nao for reconhecido ou se ``source``
            tiver tipo incompativel com o ``source_type`` informado.
        RuntimeError: Se a extracao falhar (propagado das funcoes internas).

    Exemplo::

        # PDF vindo de st.file_uploader
        uploaded = st.file_uploader("Envie um PDF", type=["pdf"])
        if uploaded:
            texto = extract_text(uploaded.read(), "pdf")

        # URL
        texto = extract_text("https://exemplo.com/pagina", "url")

        # Texto puro
        texto = extract_text("Conteudo digitado pelo usuario.", "text")
    """
    valid_types = ("pdf", "docx", "url", "text")
    source_type = source_type.strip().lower() if isinstance(source_type, str) else ""

    if source_type not in valid_types:
        raise ValueError(
            f"Tipo de fonte '{source_type}' nao reconhecido. "
            f"Valores aceitos: {', '.join(valid_types)}."
        )

    # -- Texto puro: retorno direto com strip para consistencia --
    if source_type == "text":
        if not isinstance(source, str):
            raise ValueError(
                "Para source_type='text', 'source' deve ser uma string."
            )
        result = source.strip()
        return _truncate_extracted_text(result, "texto colado")

    # -- PDF --
    if source_type == "pdf":
        if not isinstance(source, bytes):
            raise ValueError(
                "Para source_type='pdf', 'source' deve ser bytes "
                "(use file.read() do st.file_uploader)."
            )
        result = extract_from_pdf(source)
        return _truncate_extracted_text(result, "PDF")

    # -- DOCX --
    if source_type == "docx":
        if not isinstance(source, bytes):
            raise ValueError(
                "Para source_type='docx', 'source' deve ser bytes "
                "(use file.read() do st.file_uploader)."
            )
        result = extract_from_docx(source)
        return _truncate_extracted_text(result, "DOCX")

    # -- URL --
    if source_type == "url":
        if not isinstance(source, str):
            raise ValueError(
                "Para source_type='url', 'source' deve ser uma string com a URL."
            )
        result = extract_from_url(source)
        return _truncate_extracted_text(result, "URL")

    # Nunca deveria chegar aqui, mas garante seguranca
    raise ValueError(f"Tipo de fonte nao tratado: '{source_type}'.")
